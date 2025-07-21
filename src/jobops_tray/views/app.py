import re
import sys
import os
import logging
import base64
from pathlib import Path
import threading
import webbrowser
from PIL import Image
from io import BytesIO
from PySide6.QtCore import Signal, QFileSystemWatcher
import json
from jobops.utils import ResourceManager, NotificationService, check_platform_compatibility, create_desktop_entry, extract_skills, extract_skills_with_llm, build_consultant_reply_prompt
import uuid
import subprocess
from jobops.models import Document, DocumentType, JobInput
from jobops.models import Solicitation
from markdownify import markdownify
from jobops.clients import embed_structured_data


# Qt imports
try:
    from PySide6.QtWidgets import *
    from PySide6.QtCore import *
    from PySide6.QtGui import *
    _QT_AVAILABLE = True
except ImportError:
    _QT_AVAILABLE = False
    print("PySide6 is not installed. Please install PySide6.")
    sys.exit(1)


from dotenv import load_dotenv
from jobops.models import DocumentType

# Embedded base64 icon data (64x64 PNG icon)
EMBEDDED_ICON_DATA = """
iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAAOxAAADsQBlSsOGwAAABl0RVh0U29mdHdhcmUAd3d3Lmlua3NjYXBlLm9yZ5vuPBoAAAOxSURBVHic7ZtNaBNBFMefJFqrtVZbW6u01lq1Wq21aq3VWmut1lqrtdZqrdVaq7VWa63VWqu1VmuttVprtdZqrdVaq7VWa63VWqu1VmuttVprtdZqrdVaq7VWa63VWqu1VmuttVprtdZqrdVaq7VWa63VWqu1VmuttVprtdZqrdVaq7VWa60AAAD//2Q=="
"""

try:
    icon_data = base64.b64decode(EMBEDDED_ICON_DATA)
    img = Image.open(BytesIO(icon_data))
    img.show()
except Exception as e:
    print("Failed to load image:", e)

load_dotenv()

# Patch logging to always include span_id
class SpanIdLogFilter(logging.Filter):
    def filter(self, record):
        if not hasattr(record, 'span_id') or record.span_id is None:
            record.span_id = getattr(record, '_span_id', None) or str(uuid.uuid4())
        return True

logging.getLogger().addFilter(SpanIdLogFilter())

def get_span_id():
    return str(uuid.uuid4())

class JobInputDialog(QDialog):
    """Job input dialog: user provides URL (as a link) and markdown manually."""
    job_data_ready = Signal(object)
    
    def __init__(self, app_instance=None):
        super().__init__()
        # Ensure dialog appears on top when opened
        self.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        self.app_instance = app_instance
        self.setWindowTitle("Generate Motivation Letter")
        self.setFixedSize(700, 400)
        self.setWindowIcon(ResourceManager.create_app_icon())
        # Raise and activate to bring to front
        self.raise_()
        self.activateWindow()
        self._setup_ui()
        self._last_crawled_url = None
        self._last_crawled_markdown = None
        # Auto-fill URL from clipboard if an HTTP/HTTPS URL is present
        try:
            clipboard = QApplication.clipboard()
            clip_text = clipboard.text().strip()
            if re.match(r"^https?://", clip_text):
                self.url_input.setText(clip_text)
                # Trigger the URL paste handler to extract markdown automatically
                self._on_url_pasted()
                # Move focus to Generate button so Enter key will activate it
                self.generate_btn.setFocus()
        except Exception as e:
            logging.warning(f"Failed to retrieve clipboard URL: {e}")

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        # URL input (just stored as a link)
        url_layout = QHBoxLayout()
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("Paste job URL here (optional, for reference)...")
        url_layout.addWidget(QLabel("Job URL:"))
        url_layout.addWidget(self.url_input)
        layout.addLayout(url_layout)
        self.url_input.editingFinished.connect(self._on_url_pasted)
        # Company, Job Title, and Location inputs
        company_layout = QHBoxLayout()
        self.company_input = QLineEdit()
        self.company_input.setPlaceholderText("Enter company name (optional)")
        company_layout.addWidget(QLabel("Company Name:"))
        company_layout.addWidget(self.company_input)
        layout.addLayout(company_layout)

        title_layout = QHBoxLayout()
        self.title_input = QLineEdit()
        self.title_input.setPlaceholderText("Enter job title (optional)")
        title_layout.addWidget(QLabel("Job Title:"))
        title_layout.addWidget(self.title_input)
        layout.addLayout(title_layout)

        location_layout = QHBoxLayout()
        self.location_input = QLineEdit()
        self.location_input.setPlaceholderText("Enter location (optional)")
        location_layout.addWidget(QLabel("Location:"))
        location_layout.addWidget(self.location_input)
        layout.addLayout(location_layout)

        # Markdown job description (user-provided)
        layout.addWidget(QLabel("Job Description (Markdown, required):"))
        self.markdown_edit = QTextEdit()
        self.markdown_edit.setPlaceholderText("Paste or write the job description here in markdown format...")
        self.markdown_edit.setMinimumHeight(300)
        # Auto-fill company, title, location when job markdown is pasted
        self.markdown_edit.textChanged.connect(self._on_markdown_changed)
        layout.addWidget(self.markdown_edit)
        # Buttons
        button_layout = QHBoxLayout()
        self.generate_btn = QPushButton("Generate Letter")
        self.generate_btn.setDefault(True)
        self.generate_btn.setAutoDefault(True)
        self.cancel_btn = QPushButton("Cancel")
        self.generate_btn.clicked.connect(self.generate_letter)
        self.cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(self.generate_btn)
        button_layout.addWidget(self.cancel_btn)
        layout.addLayout(button_layout)

    def _on_url_pasted(self):
        url = self.url_input.text().strip()
        span_id = get_span_id()
        logging.info(f"URL pasted: {url}", extra={"span_id": span_id})
        if not url:
            return
        # Extract metadata for company, title, and location from page head
        try:
            import requests, re
            from bs4 import BeautifulSoup
            from urllib.parse import urlparse
            resp = requests.get(url, timeout=10)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.content, 'html.parser')
            # Job title
            title_meta = soup.find('meta', property='og:title') or soup.find('meta', attrs={'name':'title'})
            title = title_meta['content'] if title_meta and title_meta.get('content') else None
            if not title and soup.title and soup.title.string:
                title = soup.title.string
            if title and hasattr(self, 'title_input'):
                self.title_input.setText(title.strip())
            # Company name
            company_meta = soup.find('meta', property='og:site_name') or soup.find('meta', attrs={'name':'application-name'})
            company = company_meta['content'] if company_meta and company_meta.get('content') else None
            if not company:
                host = urlparse(url).netloc
                company = host.replace('www.', '').split('.')[0].capitalize()
            if company and hasattr(self, 'company_input'):
                self.company_input.setText(company.strip())
            # Location
            loc = None
            for attr in ('name','property'):
                m = soup.find('meta', attrs={attr: re.compile('location', re.I)})
                if m and m.get('content'):
                    loc = m['content']
                    break
            if loc and hasattr(self, 'location_input'):
                self.location_input.setText(loc.strip())
            # Auto-load job description markdown using markdownify
            try:
                # Remove scripts and styles for clean HTML
                for script in soup(["script", "style"]):
                    script.extract()
                html = str(soup)
                md = markdownify(html, heading_style="ATX")
            except Exception:
                md = soup.get_text(separator='\n', strip=True)
            if hasattr(self, 'markdown_edit'):
                self.markdown_edit.setPlainText(md)
                self._last_crawled_markdown = md
        except Exception as e:
            logging.warning(f"Metadata extraction failed: {e}", extra={"span_id": span_id})
        # Markdown auto-loaded from URL

    def generate_letter(self):
        url = self.url_input.text().strip()
        markdown = self.markdown_edit.toPlainText().strip()
        company = self.company_input.text().strip()
        title = self.title_input.text().strip()
        location = self.location_input.text().strip()
        if not markdown:
            QMessageBox.warning(self, "Error", "Job description in markdown is required.")
            return
        # Detect language of the job description
        try:
            from langdetect import detect
            detected_language = detect(markdown)
        except Exception:
            detected_language = "en"
        job_input = JobInput(
            url=url or None,
            job_markdown=markdown,
            detected_language=detected_language,
            company=company or None,
            title=title or None,
            location=location or None,
        )
        self.job_data_ready.emit(job_input)
        self.accept()

    def _on_markdown_changed(self):
        # Trigger auto-fill of company, title, and location from pasted markdown
        markdown = self.markdown_edit.toPlainText().strip()
        if not markdown:
            return
        try:
            from bs4 import BeautifulSoup
            from jobops.scrapers import WebJobScraper
            llm = getattr(self.app_instance.generator, 'llm_backend', None)
            if not llm:
                return
            # Use LLM scraper to extract job info from markdown
            soup = BeautifulSoup(markdown, 'html.parser')
            scraper = WebJobScraper(llm)
            jobinfo = scraper._extract_with_llm('', soup)
            # Populate fields if extracted
            if hasattr(self, 'company_input'):
                self.company_input.setText(jobinfo.company or '')
            if hasattr(self, 'title_input'):
                self.title_input.setText(jobinfo.title or '')
            if hasattr(self, 'location_input'):
                self.location_input.setText(jobinfo.location or '')
        except Exception:
            pass
        finally:
            # Only run once
            try:
                self.markdown_edit.textChanged.disconnect(self._on_markdown_changed)
            except Exception:
                pass

class UploadDialog(QDialog):
    """File upload dialog"""
    upload_data_ready = Signal(str, str)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Upload Document")
        self.setFixedSize(400, 200)
        self.setWindowIcon(ResourceManager.create_app_icon())
        
        self.file_path = None
        self.doc_type = None
        self.setup_ui()
    
    def setup_ui(self):
        layout = QVBoxLayout(self)
        
        # File selection
        file_layout = QHBoxLayout()
        self.file_input = QLineEdit()
        self.file_input.setPlaceholderText("Select a file...")
        self.browse_btn = QPushButton("Browse")
        self.browse_btn.clicked.connect(self.browse_file)
        
        file_layout.addWidget(self.file_input)
        file_layout.addWidget(self.browse_btn)
        layout.addLayout(file_layout)
        
        # Document type dropdown
        type_group = QGroupBox("Document Type")
        type_layout = QVBoxLayout(type_group)
        
        self.doc_type_combo = QComboBox()
        self.doc_type_combo.addItems([
            "RESUME",
            "CERTIFICATE",
            "DIPLOMA",
            "ACADEMIC TRANSCRIPT",
            "COVER_LETTER",
            "REFERENCE_LETTER",
            "PORTFOLIO",
            "WORK_SAMPLES",
            "PROJECT_DOCUMENTATION",
            "TRAINING_CERTIFICATE",
            "PROFESSIONAL_LICENSE",
            "SKILLS_ASSESSMENT",
            "PERFORMANCE_REVIEW",
            "AWARDS_AND_ACHIEVEMENTS",
            "OTHER"
        ])
        
        type_layout.addWidget(self.doc_type_combo)
        layout.addWidget(type_group)
        
        # Buttons
        button_layout = QHBoxLayout()
        self.upload_btn = QPushButton("Upload")
        self.cancel_btn = QPushButton("Cancel")
        
        self.upload_btn.clicked.connect(self.upload_document)
        self.cancel_btn.clicked.connect(self.reject)
        
        button_layout.addWidget(self.upload_btn)
        button_layout.addWidget(self.cancel_btn)
        layout.addLayout(button_layout)
    
    def browse_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, 
            "Select Document", 
            "", 
            "Documents (*.pdf *.txt *.doc *.docx);;All Files (*)"
        )
        if file_path:
            self.file_input.setText(file_path)
    
    def upload_document(self):
        file_path = self.file_input.text().strip()
        if not file_path:
            QMessageBox.warning(self, "Error", "Please select a file.")
            return
        
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "Error", "File does not exist.")
            return
        
        self.file_path = file_path
        self.doc_type = self.doc_type_combo.currentText().upper()
        self.upload_data_ready.emit(self.file_path, self.doc_type)
        self.accept()

class ConsultantInputDialog(QDialog):
    """Dialog for generating consultant reply sheet"""
    consultant_data_ready = Signal(dict)

    def __init__(self, app_instance=None):
        super().__init__()
        self.app_instance = app_instance
        self.setWindowTitle("Generate Consultant Reply Sheet")
        self.setFixedSize(700, 500)
        self.setWindowIcon(ResourceManager.create_app_icon())
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("Paste consultant email request:"))
        self.email_edit = QTextEdit()
        self.email_edit.setPlaceholderText("Paste the consultant company's email here...")
        self.email_edit.setMinimumHeight(300)
        layout.addWidget(self.email_edit)
        btn_layout = QHBoxLayout()
        generate_btn = QPushButton("Generate Answer Sheet")
        cancel_btn = QPushButton("Cancel")
        generate_btn.clicked.connect(self._on_generate)
        cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(generate_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

    def _on_generate(self):
        email_msg = self.email_edit.toPlainText().strip()
        if not email_msg:
            QMessageBox.warning(self, "Error", "Consultant email request is required.")
            return
        try:
            resume_md = self.app_instance.repository.get_latest_resume() or ""
        except Exception:
            resume_md = ""
        # Auto-detect language of the consultant request
        try:
            from langdetect import detect
            language = detect(email_msg)
        except Exception:
            language = 'en'
        self.consultant_data_ready.emit({
            'email_message': email_msg,
            'resume_markdown': resume_md,
            'language': language
        })
        self.accept()

# Add classes for parsing and input of Solicitation records
class SolicitationParseWorker(QThread):
    """Background worker to parse free-text solicitation report into structured fields"""
    finished = Signal(dict)
    error = Signal(str)

    def __init__(self, app_instance, text):
        super().__init__()
        self.app_instance = app_instance
        self.text = text

    def run(self):
        try:
            # Use the LLM backend for extraction
            generator = getattr(self.app_instance, 'generator', None)
            if generator is None:
                raise Exception('Generator not available')
            backend = getattr(generator, 'llm_backend', generator)
            # Instruct LLM to output only a valid JSON object
            prompt = (
                'You are a JSON extractor. '  
                'Return ONLY a valid JSON object with keys: datum, bedrijf, functie, status, resultaat, locatie, platform. '  
                'Do NOT include any additional text or markdown.'  
                f'\n\nReport Text:\n{self.text}'
            )
            reply = backend.generate_response(prompt)
            raw = reply.strip()
            # Extract JSON substring between first { and last }
            start = raw.find('{')
            end = raw.rfind('}')
            json_str = raw[start:end+1] if start != -1 and end != -1 else raw
            data = json.loads(json_str)
            self.finished.emit(data)
        except Exception as e:
            # Emit error for UI to display
            self.error.emit(str(e))

class SolicitationInputDialog(QDialog):
    """Dialog to create a Solicitation record, with optional auto-mapping from report text"""
    solicitation_ready = Signal(Solicitation)

    def __init__(self, app_instance=None):
        super().__init__()
        self.app_instance = app_instance
        self.setWindowTitle('Add Solicitation Record')
        self.setMinimumSize(600, 400)
        self.setWindowIcon(ResourceManager.create_app_icon())
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        # Raw text input for auto-mapping
        layout.addWidget(QLabel('Paste Sollicitatierapport content (optional for auto-mapping):'))
        self.raw_text_edit = QTextEdit()
        self.raw_text_edit.setPlaceholderText('Paste full solicitation report here...')
        self.raw_text_edit.setMinimumHeight(100)
        layout.addWidget(self.raw_text_edit)
        # Form fields
        form = QFormLayout()
        self.date_input = QLineEdit()
        self.company_input = QLineEdit()
        self.role_input = QLineEdit()
        self.status_input = QLineEdit()
        self.result_input = QLineEdit()
        self.location_input = QLineEdit()
        self.platform_input = QLineEdit()
        form.addRow('Datum:', self.date_input)
        form.addRow('Bedrijf:', self.company_input)
        form.addRow('Functie:', self.role_input)
        form.addRow('Status:', self.status_input)
        form.addRow('Resultaat:', self.result_input)
        form.addRow('Locatie:', self.location_input)
        form.addRow('Platform:', self.platform_input)
        layout.addLayout(form)
        # Buttons
        btn_layout = QHBoxLayout()
        self.parse_btn = QPushButton('Auto Map')
        self.parse_btn.clicked.connect(self._on_parse)
        self.save_btn = QPushButton('Save')
        self.save_btn.clicked.connect(self._on_save)
        self.cancel_btn = QPushButton('Cancel')
        self.cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(self.parse_btn)
        btn_layout.addWidget(self.save_btn)
        btn_layout.addWidget(self.cancel_btn)
        layout.addLayout(btn_layout)

    def _on_parse(self):
        text = self.raw_text_edit.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, 'Error', 'Please paste report content for auto-mapping.')
            return
        self.progress = QProgressDialog('Parsing report...', None, 0, 0, self)
        self.progress.setWindowTitle('JobOps')
        self.progress.setWindowModality(Qt.WindowModal)
        self.progress.setCancelButton(None)
        self.progress.show()
        self.worker = SolicitationParseWorker(self.app_instance, text)
        self.worker.finished.connect(self._on_parse_finished)
        self.worker.error.connect(self._on_parse_error)
        self.worker.start()

    def _on_parse_finished(self, data):
        self.progress.close()
        # Populate form with parsed values
        self.date_input.setText(data.get('datum',''))
        self.company_input.setText(data.get('bedrijf',''))
        self.role_input.setText(data.get('functie',''))
        self.status_input.setText(data.get('status',''))
        self.result_input.setText(data.get('resultaat',''))
        self.location_input.setText(data.get('locatie',''))
        self.platform_input.setText(data.get('platform',''))

    def _on_parse_error(self, error):
        self.progress.close()
        QMessageBox.critical(self, 'Parse Error', f'Failed to parse report: {error}')

    def _on_save(self):
        sol = Solicitation(
            datum=self.date_input.text().strip(),
            bedrijf=self.company_input.text().strip(),
            functie=self.role_input.text().strip(),
            status=self.status_input.text().strip(),
            resultaat=self.result_input.text().strip(),
            locatie=self.location_input.text().strip(),
            platform=self.platform_input.text().strip()
        )
        self.solicitation_ready.emit(sol)
        self.accept()

class SystemTrayIcon(QSystemTrayIcon):
    """Custom system tray icon with JobOps functionality"""
    
    def __init__(self, app_instance, parent=None):
        super().__init__(parent)
        self.app_instance = app_instance
        self.animation_timer = None
        self.animation_frames = self._create_animation_frames()
        self.animation_index = 0
        self.is_animating = False
        self.progress_dialog = None
        self._workers = set()  # Keep references to running workers
        self.setup_tray()
    
    def _create_animation_frames(self):
        """Create a list of QIcons for animation (spinner effect)"""
        frames = []
        base_pixmap = ResourceManager.create_app_icon().pixmap(64, 64)
        for angle in range(0, 360, 30):
            pixmap = QPixmap(base_pixmap.size())
            pixmap.fill(Qt.GlobalColor.transparent)
            painter = QPainter(pixmap)
            painter.setRenderHint(QPainter.RenderHint.Antialiasing)
            painter.translate(pixmap.width() / 2, pixmap.height() / 2)
            painter.rotate(angle)
            painter.translate(-pixmap.width() / 2, -pixmap.height() / 2)
            painter.drawPixmap(0, 0, base_pixmap)
            painter.end()
            frames.append(QIcon(pixmap))
        return frames if frames else [ResourceManager.create_app_icon()]

    def start_animation(self):
        if self.is_animating:
            return
        logging.info("Starting tray icon animation (generation in progress)")
        self.is_animating = True
        if not self.animation_timer:
            self.animation_timer = QTimer(self)
            self.animation_timer.timeout.connect(self._animate_icon)
        self.animation_index = 0
        self.animation_timer.start(100)  # 100ms per frame

    def stop_animation(self):
        if self.animation_timer:
            self.animation_timer.stop()
        self.setIcon(ResourceManager.create_app_icon())
        self.is_animating = False
        logging.info("Stopped tray icon animation (generation finished)")

    def _animate_icon(self):
        if not self.is_animating or not self.animation_frames:
            return
        self.setIcon(self.animation_frames[self.animation_index])
        self.animation_index = (self.animation_index + 1) % len(self.animation_frames)

    def setup_tray(self):
        # Set icon
        self.setIcon(ResourceManager.create_app_icon())
        self.setToolTip("JobOps - AI Motivation Letter Generator")
        
        # Create context menu
        menu = QMenu()
        
        # Add actions
        upload_action = QAction("📁 Upload", self)
        report_action = QAction("📦 Generate", self)
        reply_action = QAction("💬 Reply", self)
        investigate_action = QAction("🔍 Investigate", self)
        log_viewer_action = QAction("📝 Logs", self)
        settings_action = QAction("⚙️ Settings", self)
        help_action = QAction("❓ Help", self)
        export_action = QAction("📤 Export", self)
        quit_action = QAction("❌ Exit", self)
        
        # Connect actions
        upload_action.triggered.connect(self.upload_document)
        report_action.triggered.connect(self.generate_report)
        reply_action.triggered.connect(self.reply_to_offer)
        investigate_action.triggered.connect(self.investigate_company)
        log_viewer_action.triggered.connect(self.show_log_viewer)
        settings_action.triggered.connect(self.show_settings)
        help_action.triggered.connect(self.show_help)
        export_action.triggered.connect(self.export_document)
        quit_action.triggered.connect(self.quit_application)
        
        # Add to menu: only essential actions
        menu.addAction(upload_action)
        menu.addAction(report_action)
        menu.addAction(reply_action)
        menu.addAction(investigate_action)
        menu.addSeparator()
        menu.addAction(log_viewer_action)
        menu.addAction(settings_action)
        menu.addAction(help_action)
        menu.addSeparator()
        menu.addAction(export_action)
        menu.addAction(quit_action)
        
        self.setContextMenu(menu)
        
        # Connect signals
        self.messageClicked.connect(self.on_message_clicked)
        self.activated.connect(self.on_tray_activated)
    
    def upload_document(self):
        logging.info("User triggered: Upload Document dialog")
        dialog = UploadDialog()
        dialog.upload_data_ready.connect(self._start_upload_worker)
        dialog.exec()
    
    def _start_upload_worker(self, file_path, doc_type):
        logging.info(f"Uploading document: {file_path} as {doc_type}")
        worker = UploadWorker(self.app_instance, file_path, doc_type)
        self._workers.add(worker)
        worker.finished.connect(lambda msg, w=worker: self._on_worker_done(w, msg, is_error=False))
        worker.error.connect(lambda err, w=worker: self._on_worker_done(w, err, is_error=True))
        # Start tray animation and show progress indicator
        self.start_animation()
        self.progress_dialog = QProgressDialog("Uploading document...", None, 0, 0)
        self.progress_dialog.setWindowTitle("JobOps")
        self.progress_dialog.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.progress_dialog.setCancelButton(None)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.show()
        # Notify user via tray message
        self.showMessage("JobOps", "Uploading your document...", QSystemTrayIcon.MessageIcon.Information, 2000)
        worker.start()
    
    def generate_letter(self):
        logging.info("User triggered: Generate Letter dialog")
        dialog = JobInputDialog(self.app_instance)
        dialog.job_data_ready.connect(self._start_generate_worker)
        dialog.exec()
    
    def _start_generate_worker(self, job_input: JobInput):
        # Privacy policy consent: ask user to paste the policy text and analyze compliance
        policy_text, ok = QInputDialog.getMultiLineText(
            None,
            'Privacy Policy Consent',
            'Please paste the site\'s privacy policy text (markdown or plain text). Leave blank if not available:'
        )
        if not ok:
            return
        policy_text = policy_text.strip()
        # If policy text provided, analyze compliance via LLM
        if policy_text:
            llm = getattr(self.app_instance.generator, 'llm_backend', None)
            if llm:
                check_prompt = f"""
You are a legal expert. Analyze the following privacy policy and respond with only COMPLIANT or NON-COMPLIANT regarding GDPR and safe data handling:
{policy_text}
"""
                response = llm.generate_response(check_prompt, '')
                if 'non-compliant' in response.lower():
                    reply = QMessageBox.question(
                        None,
                        'Privacy Policy Consent',
                        'The privacy policy appears non-compliant. Do you want to continue anyway?',
                        QMessageBox.Yes | QMessageBox.No,
                        QMessageBox.No
                    )
                    if reply != QMessageBox.Yes:
                        return
        
        logging.info(f"Starting letter generation for job data: {job_input}")
        self.generate_worker = GenerateWorker(self.app_instance, job_input)
        self.generate_worker.finished.connect(self.on_generation_finished)
        self.generate_worker.error.connect(self.on_generation_error)
        self.start_animation()
        # Show progress dialog
        self.progress_dialog = QProgressDialog("Generating motivation letter...", None, 0, 0)
        self.progress_dialog.setWindowTitle("JobOps")
        self.progress_dialog.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.progress_dialog.setCancelButton(None)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.show()
        # Show notification
        self.showMessage(
            "JobOps",
            "Generating your motivation letter...",
            QSystemTrayIcon.MessageIcon.Information,
            2000
        )
        self.generate_worker.start()
    
    def show_archive(self):
        logging.info("User triggered: Show Archive dialog")
        motivations_dir = os.path.expanduser("~/.jobops/motivations")
        if sys.platform.startswith('win'):
            os.startfile(motivations_dir)
        elif sys.platform.startswith('darwin'):
            subprocess.Popen(['open', motivations_dir])
        elif sys.platform.startswith('linux'):
            subprocess.Popen(['xdg-open', motivations_dir])
        else:
            logging.error("Motivations directory does not exist.")
            QMessageBox.warning(None, "Archive", "Motivations directory does not exist.")
    
    def show_settings(self):
        logging.info("User triggered: Settings dialog")
        # Open config.json in the default text editor
        config_path = str(getattr(self.app_instance, 'config_path', Path.home() / ".jobops" / "config.json"))
        try:
            if sys.platform.startswith('win'):
                os.startfile(config_path)
            elif sys.platform.startswith('darwin'):
                subprocess.Popen(['open', config_path])
            else:
                subprocess.Popen(['xdg-open', config_path])
        except Exception as e:
            logging.error(f"Failed to open config.json: {e}")
            QMessageBox.warning(None, "Settings", f"Could not open config.json: {e}")
    
    def show_help(self):
        documentation_url = "https://github.com/codesapienbe/jobops-toolbar"
        log_message = f"User triggered: Help/documentation, opening {documentation_url}"
        logging.info(log_message)
        self.showMessage("JobOps", log_message, QSystemTrayIcon.MessageIcon.Information, 3000)
        webbrowser.open(documentation_url)
    
    def quit_application(self):
        log_message = "User triggered: Quit application"
        logging.info(log_message)
        self.showMessage("JobOps", log_message, QSystemTrayIcon.MessageIcon.Information, 3000)
        reply = QMessageBox.question(
            None, 
            "Exit JobOps", 
            "Are you sure you want to exit JobOps?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            QApplication.quit()
    
    def restart_application(self):
        log_message = "User triggered: Restart application"
        logging.info(log_message)
        self.showMessage("JobOps", log_message, QSystemTrayIcon.MessageIcon.Information, 3000)
        # make sure that the application is restarted
        exit_thread = threading.Thread(target=QApplication.quit)
        exit_thread.start()
        exit_thread.join()  # Wait for exit to complete
        restart_thread = threading.Thread(target=QApplication.exec)
        restart_thread.start()
        restart_thread.join()

    def on_upload_finished(self, message):
        log_message = f"Upload finished: {message}"
        logging.info(log_message)
        self.showMessage("JobOps", log_message, QSystemTrayIcon.MessageIcon.Information, 3000)
    
    def on_upload_error(self, error):
        log_message = f"Upload error: {error}"
        logging.error(log_message)
        self.showMessage("JobOps Error", log_message, QSystemTrayIcon.MessageIcon.Critical, 5000)
    
    def on_generation_finished(self, message):
        log_message = "Letter generation finished successfully."
        logging.info(log_message)
        self.stop_animation()
        if self.progress_dialog:
            self.progress_dialog.close()
            self.progress_dialog = None
        # Show tray notification
        self.showMessage("JobOps", log_message, QSystemTrayIcon.MessageIcon.Information, 5000)
        logging.info(log_message)
        # Fallback: use NotificationService if available
        if hasattr(self.app_instance, 'notification_service') and self.app_instance.notification_service:
            self.app_instance.notification_service.notify("JobOps", message)
            logging.info("Fallback notification service used.")
        # Show preview dialog with letter content
        self.show_letter_preview()

    def show_letter_preview(self):
        # Fetch the latest generated letter from the repository
        letters = self.app_instance.repository.get_by_type(DocumentType.COVER_LETTER)
        if not letters:
            return
        latest_letter = letters[0]  # Assuming the latest letter is the first one
        content = latest_letter.structured_content or latest_letter.raw_content
        if not content:
            return
        # Create and show preview dialog
        preview_dialog = QDialog()
        preview_dialog.setWindowTitle("Letter Preview")
        preview_dialog.setMinimumSize(800, 600)
        layout = QVBoxLayout(preview_dialog)
        text_edit = QTextEdit()
        text_edit.setPlainText(content)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)
        preview_dialog.exec()

    def on_generation_error(self, error):
        log_message = f"Letter generation error: {error}"
        logging.error(log_message)
        self.stop_animation()
        if self.progress_dialog:
            self.progress_dialog.close()
            self.progress_dialog = None
        # Show tray notification
        self.showMessage("JobOps Error", log_message, QSystemTrayIcon.MessageIcon.Critical, 5000)
        logging.info(log_message)
        # Fallback: use NotificationService if available
        if hasattr(self.app_instance, 'notification_service') and self.app_instance.notification_service:
            self.app_instance.notification_service.notify("JobOps Error", log_message)
            logging.info("Fallback notification service used.")
        # Also show error dialog
        QMessageBox.critical(None, "Generation Error", error)
    
    def on_message_clicked(self):
        log_message = "Tray message clicked."
        logging.info(log_message)
    
    def on_tray_activated(self, reason):
        if reason == QSystemTrayIcon.ActivationReason.DoubleClick:
            log_message = "Tray icon double-clicked: opening generate letter dialog."
            logging.info(log_message)
            self.generate_letter()

    def _on_worker_done(self, worker, message, is_error):
        log_message = f"Upload worker done: {message}"
        logging.info(log_message)
        self._workers.discard(worker)
        if is_error:
            self.on_upload_error(message)
        else:
            self.on_upload_finished(message)
        # Stop animation when all workers are finished
        if not self._workers:
            self.stop_animation()
            if self.progress_dialog:
                self.progress_dialog.close()
                self.progress_dialog = None

    def show_log_viewer(self):
        log_file = str(Path.home() / ".jobops" / "app.log")
        dlg = LogViewerDialog(log_file, parent=None)
        dlg.exec()

    def reply_to_offer(self):
        logging.info("User triggered: Reply dialog")
        # Prompt user for the job offer message
        message, ok = QInputDialog.getMultiLineText(None, "Reply to Message", "Paste the message to reply to (email, LinkedIn message, etc.):", "")
        if not ok or not message.strip():
            return
        # Retrieve latest resume markdown
        resume_md = getattr(self.app_instance.repository, 'get_latest_resume', lambda: None)()
        if not resume_md:
            QMessageBox.warning(None, "Error", "No resume found. Please upload your resume first.")
            return
        # Determine language from config
        config = getattr(self.app_instance, '_config', {})
        language = config.get('app_settings', {}).get('language', 'en')
        # Generate reply
        try:
            reply_text = self.app_instance.generator.generate_reply(message, resume_md, language)
        except Exception as e:
            QMessageBox.critical(None, "Error", f"Failed to generate reply: {e}")
            return
        # Show reply in a preview dialog
        dialog = QDialog()
        dialog.setWindowTitle("Reply Preview")
        dialog.setWindowIcon(ResourceManager.create_app_icon())
        dialog.resize(800, 600)
        dialog.setMinimumSize(800, 600)
        layout = QVBoxLayout(dialog)
        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        text_edit.setText(reply_text)
        layout.addWidget(text_edit)
        btn_layout = QHBoxLayout()
        copy_btn = QPushButton("Copy to Clipboard")
        copy_btn.clicked.connect(lambda: QApplication.clipboard().setText(reply_text))
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        btn_layout.addWidget(copy_btn)
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)
        dialog.exec()

    def generate_report(self):
        logging.info("User triggered: Generate Report")
        dialog = JobInputDialog(self.app_instance)
        dialog.job_data_ready.connect(self._start_report_worker)
        dialog.exec()

    def _start_report_worker(self, job_input: JobInput):
        logging.info(f"Starting report generation for job data: {job_input}")
        # Start background report worker without file export
        worker = ReportWorker(self.app_instance, job_input)
        self._workers.add(worker)
        worker.finished.connect(self.on_report_finished)
        worker.error.connect(self.on_report_error)
        self.start_animation()
        self.progress_dialog = QProgressDialog("Generating report...", None, 0, 0)
        self.progress_dialog.setWindowTitle("JobOps")
        self.progress_dialog.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.progress_dialog.setCancelButton(None)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.show()
        self.showMessage("JobOps", "Generating your report...", QSystemTrayIcon.MessageIcon.Information, 2000)
        worker.start()

    def on_report_finished(self, _save_path, _chart_path):
        log_message = "Report generation completed and saved to database."
        logging.info(log_message)
        self.stop_animation()
        if self.progress_dialog:
            self.progress_dialog.close()
            self.progress_dialog = None
        self.showMessage("JobOps", log_message, QSystemTrayIcon.MessageIcon.Information, 5000)
        # Skipping file opening since report is stored in database only

    def on_report_error(self, error):
        log_message = f"Report generation error: {error}"
        logging.error(log_message)
        self.stop_animation()
        if self.progress_dialog:
            self.progress_dialog.close()
            self.progress_dialog = None
        self.showMessage("JobOps Error", log_message, QSystemTrayIcon.MessageIcon.Critical, 5000)

    def generate_consultant_reply(self):
        """User triggers: Generate Consultant Reply dialog"""
        logging.info("User triggered: Generate Consultant Reply dialog")
        dialog = ConsultantInputDialog(self.app_instance)
        dialog.consultant_data_ready.connect(self._start_consultant_worker)
        dialog.exec()

    def _start_consultant_worker(self, data):
        """Start background worker for consultant reply sheet generation"""
        email_msg = data.get('email_message')
        resume_md = data.get('resume_markdown', '')
        language = data.get('language', 'en')
        # Start background consultant reply worker without file export
        worker = ConsultantReplyWorker(self.app_instance, email_msg, resume_md, language)
        self._workers.add(worker)
        worker.finished.connect(self.on_consultant_finished)
        worker.error.connect(self.on_consultant_error)
        self.start_animation()
        self.progress_dialog = QProgressDialog("Generating answer sheet...", None, 0, 0)
        self.progress_dialog.setWindowTitle("JobOps")
        self.progress_dialog.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.progress_dialog.setCancelButton(None)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.show()
        self.showMessage("JobOps", "Generating your answer sheet...", QSystemTrayIcon.MessageIcon.Information, 2000)
        worker.start()

    def on_consultant_finished(self, _save_path):
        """Handle successful consultant reply sheet generation"""
        log_message = "Consultant answer sheet generated and saved to database."
        logging.info(log_message)
        self.stop_animation()
        if self.progress_dialog:
            self.progress_dialog.close()
            self.progress_dialog = None
        self.showMessage("JobOps", log_message, QSystemTrayIcon.MessageIcon.Information, 5000)
        # Skipping file opening since consultant reply is stored in database only

    def on_consultant_error(self, error):
        """Handle errors during consultant reply generation"""
        log_message = f"Consultant reply generation error: {error}"
        logging.error(log_message)
        self.stop_animation()
        if self.progress_dialog:
            self.progress_dialog.close()
            self.progress_dialog = None
        self.showMessage("JobOps Error", log_message, QSystemTrayIcon.MessageIcon.Critical, 5000)
        QMessageBox.critical(None, "Consultant Reply Error", error)

    def add_solicitation_record(self):
        logging.info("User triggered: Add Solicitation Record")
        dialog = SolicitationInputDialog(self.app_instance)
        dialog.solicitation_ready.connect(self._save_solicitation)
        dialog.exec()

    def _save_solicitation(self, solicitation: Solicitation):
        try:
            rid = self.solicitation_repository.save_solicitation(solicitation)
            msg = f"Solicitation record saved (ID: {rid})"
            logging.info(msg)
            self.showMessage("JobOps", msg, QSystemTrayIcon.MessageIcon.Information, 3000)
        except Exception as e:
            logging.error(f"Error saving solicitation record: {e}")
            self.showMessage("JobOps Error", f"Failed to save solicitation: {e}", QSystemTrayIcon.MessageIcon.Critical, 5000)

    def export_document(self):
        from PySide6.QtWidgets import QInputDialog, QFileDialog, QMessageBox
        import os, zipfile
        repo = self.app_instance.repository
        group_ids = repo.list_group_ids()
        if not group_ids:
            QMessageBox.warning(None, "Export Documents", "No document groups found.")
            return
        sel_group, ok = QInputDialog.getItem(None, "Select Document Set", "Group ID:", group_ids, editable=False)
        if not ok:
            return
        docs = repo.get_by_group(sel_group)
        if not docs:
            QMessageBox.warning(None, "Export Documents", f"No documents found for group {sel_group}.")
            return

        # Ask for export format (default .md)
        formats = [".md", ".docx", ".pdf"]
        fmt, ok = QInputDialog.getItem(None, "Select Export Format", "Format:", formats, editable=False)
        if not ok:
            return

        export_dir = os.path.expanduser("~/.jobops/exports")
        os.makedirs(export_dir, exist_ok=True)
        default_name = f"documents_{sel_group}.zip"
        path, _ = QFileDialog.getSaveFileName(None, "Save Documents As", os.path.join(export_dir, default_name), "Zip (*.zip)")
        if not path:
            return
        try:
            with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                import io
                for d in docs:
                    content = d.structured_content or ""
                    if fmt == ".md":
                        filename = f"{d.type.value}_{d.id}.md"
                        data = content.encode("utf-8")
                    elif fmt == ".docx":
                        from docx import Document as DocxDocument
                        doc = DocxDocument()
                        for line in content.splitlines():
                            doc.add_paragraph(line)
                        bio = io.BytesIO()
                        doc.save(bio)
                        data = bio.getvalue()
                    elif fmt == ".pdf":
                        from reportlab.pdfgen.canvas import Canvas
                        from reportlab.lib.pagesizes import letter
                        bio = io.BytesIO()
                        c = Canvas(bio, pagesize=letter)
                        textobject = c.beginText(40, 750)
                        for line in content.splitlines():
                            textobject.textLine(line)
                        c.drawText(textobject)
                        c.save()
                        data = bio.getvalue()
                    zipf.writestr(filename, data)
            QMessageBox.information(None, "Export Documents", f"Documents zipped to {path}")
        except Exception as e:
            QMessageBox.critical(None, "Export Error", str(e))

    def investigate_company(self):
        logging.info("User triggered: Investigate Company")
        # Ask for company website URL
        website_url, ok1 = QInputDialog.getText(
            None,
            "Investigate Company",
            "Enter company website URL:",
            QLineEdit.Normal,
            ""
        )
        if not ok1 or not website_url.strip():
            return
        # Ask for company LinkedIn URL (optional)
        linkedin_url, ok2 = QInputDialog.getText(
            None,
            "Investigate Company",
            "Enter company LinkedIn URL (optional):",
            QLineEdit.Normal,
            ""
        )
        if not ok2:
            return
        website = website_url.strip()
        linkedin = linkedin_url.strip() if linkedin_url.strip() else None
        logging.info(f"Investigating website: {website}, LinkedIn: {linkedin}")
        self.investigate_worker = InvestigateWorker(self.app_instance, website, linkedin)
        self._workers.add(self.investigate_worker)
        self.investigate_worker.finished.connect(self.on_investigation_finished)
        self.investigate_worker.error.connect(self.on_investigation_error)
        self.start_animation()
        self.progress_dialog = QProgressDialog("Investigating company...", None, 0, 0)
        self.progress_dialog.setWindowTitle("JobOps")
        self.progress_dialog.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.progress_dialog.setCancelButton(None)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.show()
        self.investigate_worker.start()

    def on_investigation_finished(self, message):
        logging.info("Company investigation finished")
        self.stop_animation()
        if self.progress_dialog:
            self.progress_dialog.close()
            self.progress_dialog = None
        self.showMessage("Company Investigation", "Investigation completed", QSystemTrayIcon.MessageIcon.Information, 5000)
        dialog = QDialog()
        dialog.setWindowTitle("Company Investigation Result")
        dialog.setMinimumSize(600, 400)
        layout = QVBoxLayout(dialog)
        text_edit = QTextEdit()
        text_edit.setPlainText(message)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)
        dialog.exec()

    def on_investigation_error(self, error):
        logging.error(f"Company investigation error: {error}")
        self.stop_animation()
        if self.progress_dialog:
            self.progress_dialog.close()
            self.progress_dialog = None
        self.showMessage("Company Investigation Error", error, QSystemTrayIcon.MessageIcon.Critical, 5000)
        QMessageBox.critical(None, "Investigation Error", error)

class UploadWorker(QThread):
    """Background worker for document upload"""
    finished = Signal(str)
    error = Signal(str)
    
    def __init__(self, app_instance, file_path, doc_type):
        super().__init__()
        self.app_instance = app_instance
        self.file_path = file_path
        self.doc_type = doc_type
    
    def run(self):
        try:
            import os
            import shutil
            from datetime import datetime
            from markitdown import MarkItDown
            from openai import OpenAI
            from jobops.models import Document, DocumentType
            repository = getattr(self.app_instance, 'repository', None)
            span_id = get_span_id()
            logging.info(f"UploadWorker started for file: {self.file_path}, doc_type: {self.doc_type}", extra={"span_id": span_id})
            if repository is None:
                logging.error("Application is missing repository.", extra={"span_id": span_id})
                raise Exception("Application is missing repository.")

            # Copy uploaded file to ~/.jobops/resumes/
            resumes_dir = os.path.expanduser("~/.jobops/resumes")
            os.makedirs(resumes_dir, exist_ok=True)
            dest_path = os.path.join(resumes_dir, os.path.basename(self.file_path))
            shutil.copy2(self.file_path, dest_path)
            logging.info(f"Copied file to: {dest_path}", extra={"span_id": span_id})

            # Use dest_path as the filename for the Document
            filename_for_db = dest_path

            # Determine file extension
            _, ext = os.path.splitext(self.file_path)
            ext = ext.lower()

            if ext == ".md":
                # If markdown, just read the content
                with open(self.file_path, "r", encoding="utf-8") as f:
                    structured_content = f.read()
                raw_content = structured_content
                logging.info(f"Markdown file detected, skipping MarkItDown conversion.", extra={"span_id": span_id})
            else:
                # Get the LLM backend and model
                generator = getattr(self.app_instance, 'generator', None)
                llm_backend = getattr(generator, 'llm_backend', None)
                logging.info(f"Detected backend: {llm_backend}", extra={"span_id": span_id})
                if not llm_backend:
                    logging.error("No LLM backend available for MarkItDown.", extra={"span_id": span_id})
                    raise Exception("No LLM backend available for MarkItDown.")
                backend_name = getattr(llm_backend, 'name', '').lower()
                llm_model = getattr(llm_backend, 'model', 'gpt-4o')
                logging.info(f"Backend name: {backend_name}, model: {llm_model}", extra={"span_id": span_id})
                if backend_name in ("openai", "groq"):
                    llm_client = llm_backend.client
                    logging.info(f"Using OpenAI/Groq client: {type(llm_client)}", extra={"span_id": span_id})
                elif backend_name == "ollama":
                    base_url = getattr(llm_backend, 'base_url', 'http://localhost:11434')
                    logging.info(f"Using Ollama backend, base_url: {base_url}", extra={"span_id": span_id})
                    llm_client = OpenAI(base_url=f"{base_url}/v1", api_key="ollama")
                else:
                    logging.error(f"MarkItDown integration is not yet supported for backend: {backend_name}", extra={"span_id": span_id})
                    raise Exception(f"MarkItDown integration is not yet supported for backend: {backend_name}")
                md = MarkItDown(llm_client=llm_client, llm_model=llm_model)
                logging.info(f"Starting MarkItDown extraction for file: {self.file_path}", extra={"span_id": span_id})
                try:
                    result = md.convert(self.file_path)
                except Exception as e:
                    # Check for MissingDependencyException in the error message
                    if "MissingDependencyException" in str(e) and ("docx" in ext or "docx" in str(e)):
                        msg = (
                            "File conversion failed: MarkItDown requires extra dependencies to read .docx files. "
                            "Please install with: pip install markitdown[docx] or pip install markitdown[all]"
                        )
                        logging.error(msg, extra={"span_id": span_id})
                        self.error.emit(msg)
                        return
                    else:
                        raise
                structured_content = result.text_content
                raw_content = result.raw_text if hasattr(result, 'raw_text') else structured_content
                logging.info(f"Extraction complete, structured length: {len(structured_content)}", extra={"span_id": span_id})

            # Compute embedding for RAG similarity
            embedding = embed_structured_data(structured_content)

            doc_type_enum = DocumentType.RESUME if self.doc_type.upper() == "RESUME" else DocumentType.CERTIFICATION
            doc = Document(
                type=doc_type_enum,
                raw_content=raw_content,
                structured_content=structured_content,
                uploaded_at=datetime.now(),
                embedding=embedding,
                group_id=str(uuid.uuid4())
            )
            repository.save(doc)
            logging.info(f"Document saved to database: {filename_for_db}", extra={"span_id": span_id})
            self.finished.emit(f"Document uploaded and parsed successfully: {os.path.basename(self.file_path)}")
        except Exception as e:
            import traceback
            logging.error(f"Upload failed: {str(e)}\n{traceback.format_exc()}", extra={"span_id": span_id})
            self.error.emit(f"Upload failed: {str(e)}\n{traceback.format_exc()}")

class GenerateWorker(QThread):
    """Background worker for letter generation"""
    finished = Signal(str)
    error = Signal(str)
    
    def __init__(self, app_instance, job_data: JobInput):
        super().__init__()
        self.app_instance = app_instance
        self.job_data = job_data
    
    def run(self):
        try:
            log_message = "GenerateWorker started."
            logging.info(log_message)
            repository = getattr(self.app_instance, 'repository', None)
            generator = getattr(self.app_instance, 'generator', None)
            if repository is None or generator is None:
                log_message = "Application is missing repository or generator."
                logging.error(log_message)
                raise Exception("Application is missing repository or generator.")
            resume_markdown = repository.get_latest_resume()
            if resume_markdown is None:
                log_message = "No resume found. Please upload your resume first."
                logging.error(log_message)
                raise Exception("No resume found. Please upload your resume first.")
            logging.info(f"Resume markdown (first 100 chars): {resume_markdown[:100]}")
            job_markdown = self.job_data.job_markdown
            if not job_markdown:
                raise Exception("Job description markdown is required.")
            detected_language = self.job_data.detected_language
            url = self.job_data.url or ''
            # Get truncation config from app_instance if available
            config = getattr(self.app_instance, '_config', None)
            backend_type = config.get('backend', 'ollama') if config else None
            backend_settings = config.get('backend_settings', {}) if config else {}
            model_conf = backend_settings.get(backend_type, {}) if backend_settings else {}
            trunc_config = model_conf if model_conf else {}

            # ------------------------------------------------------------------
            # 1) Generate Cover Letter
            # ------------------------------------------------------------------
            letter = generator.generate_from_markdown(
                job_markdown,
                resume_markdown,
                detected_language,
                url=url,
                config=trunc_config,
                company_name=self.job_data.company or '',
                job_title=self.job_data.title or '',
                location=self.job_data.location or ''
            )

            # ------------------------------------------------------------------
            # 2) Generate Tailored Resume and Skills Report
            # ------------------------------------------------------------------
            tailored_resume = generator.generate_optimized_resume_from_markdown(
                job_markdown,
                resume_markdown,
                detected_language,
                requirements=self.job_data.requirements or '',
                config=trunc_config,
            )

            # Extract skills
            from jobops.utils import extract_skills, extract_skills_with_llm
            llm_backend = getattr(self.app_instance.generator, 'llm_backend', None)
            job_requirements = self.job_data.requirements or ''
            skill_data = (
                extract_skills_with_llm(llm_backend, resume_markdown, job_markdown + "\n" + job_requirements)
                if llm_backend
                else None
            )
            if skill_data:
                matched = sorted(set(skill_data['matching_skills']))
                missing = sorted(set(skill_data['missing_skills']))
                extra = sorted(set(skill_data['extra_skills']))
            else:
                resume_skills = extract_skills(resume_markdown)
                job_skills = extract_skills(job_markdown + "\n" + job_requirements)
                matched = sorted(resume_skills & job_skills)
                missing = sorted(job_skills - resume_skills)
                extra = sorted(resume_skills - job_skills)

            total_required = len(matched) + len(missing)
            summary = f"You have matched {len(matched)} of {total_required} required skills."
            report_lines = [
                "# Skills Match Report",
                "",
                summary,
                "",
                "## Skill Details",
                "",
            ]
            if matched:
                report_lines.append("**Matched Skills:**")
                report_lines.extend([f"- {s}" for s in matched])
            if missing:
                report_lines.append("")
                report_lines.append("**Missing Skills:**")
                report_lines.extend([f"- {s}" for s in missing])
            if extra:
                report_lines.append("")
                report_lines.append("**Additional Skills:**")
                report_lines.extend([f"- {s}" for s in extra])
            report_md = "\n".join(report_lines)
            # Assign group_id for this document set
            group_id = str(uuid.uuid4())
            # Store job description
            doc_id_job = str(uuid.uuid4())
            doc_job = Document(
                id=doc_id_job,
                type=DocumentType.JOB_DESCRIPTION,
                raw_content=job_markdown,
                structured_content=job_markdown,
                group_id=group_id
            )
            repository.save(doc_job)
            # Store tailored resume
            doc_id_resume = str(uuid.uuid4())
            doc_resume = Document(
                id=doc_id_resume,
                type=DocumentType.RESUME,
                raw_content=tailored_resume,
                structured_content=tailored_resume,
                group_id=group_id
            )
            repository.save(doc_resume)
            # Store cover letter
            doc_id_letter = str(uuid.uuid4())
            doc_letter = Document(
                id=doc_id_letter,
                type=DocumentType.COVER_LETTER,
                raw_content=letter.content,
                structured_content=letter.content,
                group_id=group_id
            )
            repository.save(doc_letter)
            # Store match report
            doc_id_report = str(uuid.uuid4())
            doc_report = Document(
                id=doc_id_report,
                type=DocumentType.OTHER,
                raw_content=report_md,
                structured_content=report_md,
                group_id=group_id
            )
            repository.save(doc_report)
            self.finished.emit("Motivation package generated and saved to database.")
        except Exception as e:
            self.error.emit(str(e))

# ---------------------------------------------------------------------------
# ReportWorker (restored after refactor)
# ---------------------------------------------------------------------------


class ReportWorker(QThread):
    """Background worker for generating tailored resume, cover letter, and skills match report."""

    finished = Signal(str, str)
    error = Signal(str)

    def __init__(self, app_instance, job_data: JobInput):
        super().__init__()
        self.app_instance = app_instance
        self.job_data = job_data

    def run(self):
        try:
            repository = getattr(self.app_instance, 'repository', None)
            generator = getattr(self.app_instance, 'generator', None)
            if repository is None or generator is None:
                raise Exception("Application is missing repository or generator.")

            resume_markdown = repository.get_latest_resume()
            if resume_markdown is None:
                raise Exception("No resume found. Please upload your resume first.")

            job_markdown = self.job_data.job_markdown
            if not job_markdown:
                raise Exception("Job description markdown is required.")

            detected_language = self.job_data.detected_language
            url = self.job_data.url or ''

            config = getattr(self.app_instance, '_config', None)
            backend_type = config.get('backend', 'ollama') if config else None
            backend_settings = config.get('backend_settings', {}) if config else {}
            model_conf = backend_settings.get(backend_type, {}) if backend_settings else {}
            trunc_config = model_conf if model_conf else {}

            letter = generator.generate_from_markdown(
                job_markdown,
                resume_markdown,
                detected_language,
                url=url,
                config=trunc_config,
            )

            tailored_resume = generator.generate_optimized_resume_from_markdown(
                job_markdown,
                resume_markdown,
                detected_language,
                requirements=self.job_data.requirements or '',
                config=trunc_config,
            )

            from jobops.utils import extract_skills, extract_skills_with_llm

            job_requirements = self.job_data.requirements or ''
            llm_backend = getattr(self.app_instance.generator, 'llm_backend', None)
            skill_data = (
                extract_skills_with_llm(llm_backend, resume_markdown, job_markdown + "\n" + job_requirements)
                if llm_backend else None
            )

            if skill_data:
                matched = sorted(set(skill_data['matching_skills']))
                missing = sorted(set(skill_data['missing_skills']))
                extra = sorted(set(skill_data['extra_skills']))
            else:
                resume_skills = extract_skills(resume_markdown)
                job_skills = extract_skills(job_markdown + "\n" + job_requirements)
                matched = sorted(resume_skills & job_skills)
                missing = sorted(job_skills - resume_skills)
                extra = sorted(resume_skills - job_skills)

            total_required = len(matched) + len(missing)
            summary = f"You have matched {len(matched)} of {total_required} required skills."
            report_lines = [
                "# Skills Match Report",
                "",
                summary,
                "",
                "## Skill Details",
                "",
            ]
            if matched:
                report_lines.append("**Matched Skills:**")
                report_lines.extend([f"- {s}" for s in matched])
            if missing:
                report_lines.append("")
                report_lines.append("**Missing Skills:**")
                report_lines.extend([f"- {s}" for s in missing])
            if extra:
                report_lines.append("")
                report_lines.append("**Additional Skills:**")
                report_lines.extend([f"- {s}" for s in extra])
            report_md = "\n".join(report_lines)

            group_id = str(uuid.uuid4())

            from jobops.models import Document, DocumentType

            for doc_type, content in [
                (DocumentType.JOB_DESCRIPTION, job_markdown),
                (DocumentType.RESUME, tailored_resume),
                (DocumentType.COVER_LETTER, letter.content),
                (DocumentType.OTHER, report_md),
            ]:
                repository.save(
                    Document(
                        id=str(uuid.uuid4()),
                        type=doc_type,
                        raw_content=content,
                        structured_content=content,
                        group_id=group_id,
                    )
                )

            self.finished.emit("", "")
        except Exception as e:
            self.error.emit(str(e))

class ConsultantReplyWorker(QThread):
    """Background worker for consultant reply sheet generation"""
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, app_instance, email_message, resume_markdown, language):
        super().__init__()
        self.app_instance = app_instance
        self.email_message = email_message
        self.resume_markdown = resume_markdown
        self.language = language

    def run(self):
        try:
            generator = getattr(self.app_instance, 'generator', None)
            if generator is None:
                raise Exception("Application generator is missing.")
            llm_backend = getattr(generator, 'llm_backend', None)
            # Build prompt and generate reply
            prompt = build_consultant_reply_prompt(self.email_message, self.resume_markdown, self.language)
            if llm_backend:
                reply_md = llm_backend.generate_response(prompt)
            else:
                reply_md = generator.backend.generate_response(prompt, "")
            # Store consultant reply only in the database, remove file export
            repository = getattr(self.app_instance, 'repository', None)
            if repository is None:
                raise Exception("Application is missing repository.")
            doc_id = str(uuid.uuid4())
            doc = Document(
                id=doc_id,
                type=DocumentType.OTHER,
                raw_content=reply_md,
                structured_content=reply_md
            )
            repository.save(doc)
            self.finished.emit("")
        except Exception as e:
            self.error.emit(str(e))

class JobOpsQtApplication(QApplication):
    """Main Qt application class"""
    
    def __init__(self, args):
        super().__init__(args)
        
        # Set application properties
        self.setApplicationName("JobOps")
        self.setApplicationVersion("2.0.0")
        self.setOrganizationName("JobOps")
        self.setQuitOnLastWindowClosed(False)  # Keep running in system tray
        
        # Initialize application data
        self.base_dir = Path.home() / ".jobops"
        self.base_dir.mkdir(exist_ok=True)

        # --- Load config and initialize repository and generator ---
        from jobops.repositories import SQLiteDocumentRepository
        db_path = str(self.base_dir / "jobops.db")
        self.repository = SQLiteDocumentRepository(db_path)

        # Load config.json
        self.config_path = self.base_dir / "config.json"
        self._init_config_and_generator()
        # ---------------------------------------------------
        
        # Output format and debug level from config
        self.output_format = self._config.get("output_format", "markdown").lower()
        self.debug = self._config.get("debug", False)
        self.setup_logging()
        self.notification_service = NotificationService()
        self.system_tray = None
        
        self.setup_system_tray()
        # Automatically prompt for resume upload if none exists
        if self.repository.get_latest_resume() is None:
            self.system_tray.upload_document()
        # Add config.json watcher
        self.config_watcher = QFileSystemWatcher([str(self.config_path)])
        self.config_watcher.fileChanged.connect(self.on_config_changed)
    
    def _init_config_and_generator(self):
        """Load config and initialize generator and settings."""
        if self.config_path.exists():
            with open(self.config_path, "r", encoding="utf-8") as f:
                config = json.load(f)
        else:
            config = {}
        self._config = config
        backend_type = config.get("backend", "ollama")
        backend_settings = config.get("backend_settings", {})
        app_settings = config.get("app_settings", {})
        tokens = app_settings.get("tokens", {})
        backend_conf = backend_settings.get(backend_type, {})
        from jobops.utils import ConcreteLetterGenerator
        from jobops.clients import LLMBackendFactory
        backend = LLMBackendFactory.create(backend_type, backend_conf, tokens)
        self.generator = ConcreteLetterGenerator(backend)
        # Output format and debug level from config
        self.output_format = app_settings.get("output_format", "markdown").lower()
        self.debug = config.get("debug", False)
        # Health check for model validity
        try:
            if hasattr(backend, 'health_check'):
                ok = backend.health_check()
                if not ok:
                    # Provide model list URL based on backend
                    model_urls = {
                        'groq': 'https://console.groq.com/docs/models',
                        'openai': 'https://platform.openai.com/docs/models',
                        'ollama': 'https://ollama.com/library',
                        'gemini': 'https://ai.google.dev/models',
                    }
                    url = model_urls.get(backend_type, None)
                    msg = f"The selected model for backend '{backend_type}' is not available or invalid."
                    if url:
                        msg += f"\nCheck available models here: {url}"
                        import webbrowser
                        webbrowser.open(url)
                    QMessageBox.warning(None, "Model Health Check Failed", msg)
        except Exception as e:
            logging.warning(f"Model health check failed: {e}")

    def on_config_changed(self, path):
        # QFileSystemWatcher sometimes emits multiple times, so reload defensively
        try:
            self._init_config_and_generator()
            self.setup_logging()
            self.notification_service.notify(
                "JobOps Config Reloaded",
                "Settings have been reloaded from config.json."
            )
            log_message = "Config reloaded from config.json"
            logging.info(log_message)
        except Exception as e:
            log_message = f"Failed to reload config: {e}"
            logging.error(log_message)

    def setup_logging(self):
        """Setup application logging"""
        log_file = self.base_dir / 'app.log'
        log_level = logging.DEBUG if getattr(self, 'debug', False) else logging.INFO
        try:
            from pythonjsonlogger import jsonlogger
            file_handler = logging.FileHandler(log_file, encoding='utf-8')
            json_format = '%(asctime)s %(levelname)s %(name)s %(message)s %(span_id)s %(trace_id)s'
            file_handler.setFormatter(jsonlogger.JsonFormatter(json_format))
        except ImportError:
            # Fallback: custom JSON formatter
            class SimpleJsonFormatter(logging.Formatter):
                def format(self, record):
                    import json
                    log_record = {
                        'timestamp': self.formatTime(record, self.datefmt),
                        'level': record.levelname,
                        'logger': record.name,
                        'message': record.getMessage(),
                        'span_id': getattr(record, 'span_id', None),
                        'trace_id': getattr(record, 'trace_id', None),
                    }
                    return json.dumps(log_record)
            file_handler = logging.FileHandler(log_file, encoding='utf-8')
            file_handler.setFormatter(SimpleJsonFormatter())
        stream_handler = logging.StreamHandler()
        stream_handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(name)s: %(message)s'))
        root_logger = logging.getLogger()
        root_logger.setLevel(log_level)
        # Remove all handlers first (avoid duplicate logs)
        for h in root_logger.handlers[:]:
            root_logger.removeHandler(h)
        root_logger.addHandler(file_handler)
        root_logger.addHandler(stream_handler)
    
    def setup_system_tray(self):
        """Initialize system tray"""
        if not QSystemTrayIcon.isSystemTrayAvailable():
            QMessageBox.critical(
                None,
                "System Tray",
                "System tray is not available on this system."
            )
            sys.exit(1)
        
        self.system_tray = SystemTrayIcon(self)
        self.notification_service.set_system_tray(self.system_tray)
        self.system_tray.show()
        
        # Show startup notification
        self.system_tray.showMessage(
            "JobOps Started",
            "JobOps is running in the system tray. Right-click the icon to access features.",
            QSystemTrayIcon.MessageIcon.Information,
            3000
        )

def main():
    """Main application entry point"""
    
    # Check platform compatibility
    if not check_platform_compatibility():
        sys.exit(1)
    
    # Create Qt application
    app = JobOpsQtApplication(sys.argv)
    
    # Setup platform-specific features
    if sys.platform.startswith('linux'):
        create_desktop_entry()
    
    # Setup global hotkey: Win+Shift+J on Windows, application shortcut on others
    if sys.platform.startswith("win"):
        try:
            import ctypes
            from ctypes import wintypes
            WM_HOTKEY = 0x0312
            MOD_WIN = 0x0008
            MOD_SHIFT = 0x0004
            HOTKEY_ID = 1
            class WindowsHotkeyFilter(QAbstractNativeEventFilter):
                def __init__(self, callback):
                    super().__init__()
                    self.callback = callback
                def nativeEventFilter(self, _eventType, message):
                    msg = ctypes.wintypes.MSG.from_address(int(message.__int__()))
                    if msg.message == WM_HOTKEY and msg.wParam == HOTKEY_ID:
                        QTimer.singleShot(0, self.callback)
                    return False, 0
            user32 = ctypes.windll.user32
            if not user32.RegisterHotKey(None, HOTKEY_ID, MOD_WIN | MOD_SHIFT, ord("J")):
                raise Exception("RegisterHotKey failed")
            hotkey_filter = WindowsHotkeyFilter(app.system_tray.generate_letter)
            app.installNativeEventFilter(hotkey_filter)
            logging.info("Global hotkey registered: Win+Shift+J")
        except Exception as e:
            logging.warning(f"Failed to register global hotkey: {e}")
    else:
        try:
            # Application-level shortcut for quick letter generation (Ctrl+Alt+J)
            shortcut = QShortcut(QKeySequence("Ctrl+Alt+J"), None)
            shortcut.setContext(Qt.ApplicationShortcut)
            shortcut.activated.connect(app.system_tray.generate_letter)
            logging.info("Application shortcut registered: Ctrl+Alt+J")
        except Exception as e:
            logging.warning(f"Failed to setup application shortcut: {e}")
    
    # Show startup message
    log_message = "JobOps Qt application started successfully"
    logging.info(log_message)
    
    # Write PID file for monitoring
    config_dir = os.path.expanduser("~/.jobops")
    os.makedirs(config_dir, exist_ok=True)
    pid_file = os.path.join(config_dir, "jobops_tray_pid.txt")
    with open(pid_file, "w") as f:
        f.write(str(os.getpid()))
    try:
        # Start the application event loop
        try:
            sys.exit(app.exec())
        except KeyboardInterrupt:
            log_message = "Application interrupted by user"
            logging.info(log_message)
            sys.exit(0)
        except Exception as e:
            log_message = f"Application error: {e}"
            logging.error(log_message)
            # Show error dialog
            QMessageBox.critical(
                None,
                "JobOps Error",
                f"An unexpected error occurred:\n\n{str(e)}\n\nPlease check the log file for details."
            )
            sys.exit(1)
    finally:
        # Remove PID file on exit
        try:
            os.remove(pid_file)
        except Exception:
            pass

if __name__ == "__main__":
    main()

# ---------------------------------------------------------------------------
# LogViewerDialog (restored)
# ---------------------------------------------------------------------------


class LogViewerDialog(QDialog):
    def __init__(self, log_file, parent=None):
        super().__init__(parent)
        self.setWindowTitle("JobOps Log Viewer")
        self.setMinimumSize(900, 600)
        layout = QVBoxLayout(self)
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search logs...")
        self.search_input.textChanged.connect(self.filter_logs)
        layout.addWidget(self.search_input)

        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels([
            "Timestamp",
            "Level",
            "Logger",
            "Message",
            "span_id",
            "trace_id",
        ])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)

        self.log_file = log_file
        self.all_logs: list[dict] = []
        self.load_logs()

    # -------------------------- Helpers ----------------------------------

    def load_logs(self):
        self.all_logs.clear()
        self.table.setRowCount(0)
        if not os.path.exists(self.log_file):
            return
        import json as _json
        with open(self.log_file, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    log = _json.loads(line)
                except Exception:
                    # Fallback: treat as plain text
                    log = {
                        "timestamp": "",
                        "level": "",
                        "logger": "",
                        "message": line,
                        "span_id": "",
                        "trace_id": "",
                    }
                self.all_logs.append(log)
        self.display_logs(self.all_logs)

    def display_logs(self, logs):
        self.table.setRowCount(len(logs))
        from PySide6.QtGui import QColor
        for row, log in enumerate(logs):
            ts = log.get("timestamp", log.get("asctime", ""))
            lvl = log.get("level", log.get("levelname", ""))
            logger = log.get("logger", log.get("name", ""))
            msg = log.get("message", "")
            span_id = log.get("span_id", "")
            trace_id = log.get("trace_id", "")
            for col, val in enumerate([ts, lvl, logger, msg, span_id, trace_id]):
                item = QTableWidgetItem(str(val))
                if col == 1:  # Level column
                    if lvl == "ERROR":
                        item.setForeground(QColor("red"))
                    elif lvl == "WARNING":
                        item.setForeground(QColor("orange"))
                    elif lvl == "INFO":
                        item.setForeground(QColor("blue"))
                    elif lvl == "DEBUG":
                        item.setForeground(QColor("gray"))
                self.table.setItem(row, col, item)

    def filter_logs(self, text):
        text = text.lower()
        import json as _json
        filtered = [log for log in self.all_logs if text in _json.dumps(log).lower()]
        self.display_logs(filtered)


# ---------------------------------------------------------------------------
# InvestigateWorker (restored)
# ---------------------------------------------------------------------------


class InvestigateWorker(QThread):
    """Background worker for company investigation"""

    finished = Signal(str)
    error = Signal(str)

    def __init__(self, app_instance, website_url, linkedin_url=None):
        super().__init__()
        self.app_instance = app_instance
        self.website_url = website_url
        self.linkedin_url = linkedin_url

    def run(self):
        try:
            logging.info(
                "InvestigateWorker started for Website: %s, LinkedIn: %s",
                self.website_url,
                self.linkedin_url,
            )
            import asyncio
            from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig, CacheMode

            # Configure crawler
            browser_conf = BrowserConfig(headless=True, verbose=False)
            run_conf = CrawlerRunConfig(cache_mode=CacheMode.ENABLED)

            async def crawl_site(url):
                async with AsyncWebCrawler(config=browser_conf) as crawler:
                    return await crawler.arun(url=url, config=run_conf)

            web_result = asyncio.run(crawl_site(self.website_url))
            website_content = getattr(web_result.markdown, "raw_markdown", str(web_result.markdown))

            linkedin_content = ""
            if self.linkedin_url:
                li_result = asyncio.run(crawl_site(self.linkedin_url))
                linkedin_content = getattr(li_result.markdown, "raw_markdown", str(li_result.markdown))

            combined_content = f"Website Content:\n{website_content}\n"
            if linkedin_content:
                combined_content += f"\nLinkedIn Content:\n{linkedin_content}\n"

            llm_backend = getattr(self.app_instance.generator, "llm_backend", None)
            if not llm_backend:
                raise Exception("No LLM backend available for analysis.")

            analysis_prompt = f"""
You are an expert in corporate intelligence and security.
Using the following information, provide:
1) summary: Brief summary of the company.
2) employees_count: Number of employees based on LinkedIn data.
3) followers_count: Number of LinkedIn followers.
4) org_structure: List the main departments or organizational units.
5) red_flags: Any potential red flags or suspicious aspects.
6) rating: Trustworthiness on a scale from 1 (not trustworthy) to 5 (very trustworthy).
Output only JSON with keys: summary (string), employees_count (integer), followers_count (integer), org_structure (list of strings), red_flags (list of strings), rating (integer).

{combined_content}
"""
            analysis = llm_backend.generate_response(analysis_prompt.strip(), "")
            self.finished.emit(analysis)
        except Exception as e:
            logging.error("InvestigateWorker error: %s", e)
            self.error.emit(str(e))
